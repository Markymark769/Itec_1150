import requests
from openpyxl import Workbook
import docx
# grab World Bank data
url = 'http://api.worldbank.org/v2/country?format=json&per_page=400'
world_bank_data = requests.get(url).json()
# use World Bank data to extract countries and their capitals
country_info = world_bank_data[1]
countries = {}
for data in country_info:
    country = data['name']
    capital = data['capitalCity']
    if capital:
        countries[country] = capital

# setup spreadsheet stuff
workbook = Workbook()
worksheet = workbook.active
worksheet.title = 'World Countries'
worksheet.cell(1, 1, 'Country')
worksheet.cell(1, 2, 'Capital')
worksheet.column_dimensions['A'].width = 40     # width for column A
worksheet.column_dimensions['B'].width = 50     # width for column B


# setup document stuff
document = docx.Document()
document.add_paragraph('Countries and their Capital Cities', 'Title')
# loop over all the data
row = 2 # skip row 1 for the headers
for country, capital in countries.items():
    # add to the spreadsheet
    worksheet.cell(row, 1, country)
    worksheet.cell(row, 2, capital)
    row += 1
    # add to the document
    document.add_paragraph(f'The capital of {country} is {capital}.')

# use similar filenames for both
filename = 'Country Capitals'
sheet_filename = filename + '.xlsx'
doc_filename = filename + '.docx'
# save the spreadsheet
workbook.save(sheet_filename)
# save the document
document.save(doc_filename)