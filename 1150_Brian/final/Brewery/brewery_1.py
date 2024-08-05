import docx
import requests
import random

# Step 1: Make the API call
url = 'https://api.openbrewerydb.org/v1/breweries'
response = requests.get(url)

# Check if the request was successful
if response.status_code == 200:
    response_list_of_all_breweries = response.json()
else:
    print("Failed to retrieve data")
    response_list_of_all_breweries = []

# Step 2: Select 20 random breweries from the list
if response_list_of_all_breweries:
    list_of_random_breweries_chosen = random.sample(response_list_of_all_breweries, 20)
else:
    list_of_random_breweries_chosen = []

# Step 3: Create a Word document
brewery_word_document = docx.Document()

# Step 4: Iterate over the selected breweries and get details
for one_brewery in list_of_random_breweries_chosen:
    brewery_id = one_brewery['id']
    url_for_one_brewery_details = f'https://api.openbrewerydb.org/v1/breweries/{brewery_id}'

    # Get details of each brewery
    one_brewery_details_response = requests.get(url_for_one_brewery_details)

    if one_brewery_details_response.status_code == 200:
        one_brewery_details = one_brewery_details_response.json()

        # Extract and print brewery details
        brewery_name = one_brewery_details.get('name', 'N/A')
        brewery_website = one_brewery_details.get('website_url', 'N/A')

        # Add brewery name to the document
        brewery_word_document.add_paragraph(brewery_name, 'Title')

        # Add website URL to the document
        brewery_word_document.add_paragraph('Website', 'Heading 1')
        brewery_word_document.add_paragraph(brewery_website)

    else:
        print(f"Failed to retrieve details for brewery ID: {brewery_id}")

# Step 6: Save the Word document
brewery_word_document.save('Brewery_Guide.docx')
