"""
Mark Porraz
8/6/2024
Brewery_API.py

Here is where I make a call to pull 20 randomly selected breweries and place them in a Word document

Here is the grand website where data is pulled from:
https://www.openbrewerydb.org/documentation#by_dist
"""

import docx
import requests
import random

# step 1 make API call
url = 'https://api.openbrewerydb.org/v1/breweries'
response_list_of_all_breweries = requests.get(url).json()
# print(brewery_list_url), no longer needed --- used for checking


# step 2: select from a list of 20 randomly selected breweries
list_of_random_breweries_chosen = random.sample(response_list_of_all_breweries, 20)
# print(list_of_random_breweries_chosen) --- used for checking

# step 3: creating a Word document
brewery_word_document = docx.Document()

# step 4: create a for loop among the chosen breweries and get the details for the Word document
for one_brewery in list_of_random_breweries_chosen:
    # print(one_brewery) --- used for checking
    brewery_id = one_brewery['id']

    # print(brewery_id) --- used for checking
    url_for_one_brewery_details = f'https://api.openbrewerydb.org/v1/breweries/{brewery_id}'
    # print(url_for_one_brewery_details) --- used for checking
    one_brewery_details_response = requests.get(url_for_one_brewery_details)
    # print(one_brewery_details_response) --- used for checking

    if one_brewery_details_response.status_code == 200:  # if status code is 200, which means that request has succeeded
        one_brewery_details = one_brewery_details_response.json()  # get the details response from the chosen ['id']
        # from one brewery

        # extracting details from the brewery
        # from the API, pull the information from the following keys from the dictionary to get value
        brewery_name = one_brewery_details.get('name')
        brewery_address = one_brewery_details.get('address_1')
        brewery_city = one_brewery_details.get('city')
        brewery_state = one_brewery_details.get('state')
        brewery_type = one_brewery_details.get('brewery_type')
        brewery_website = one_brewery_details.get('website_url')

# formatting and adding the details to the Word document

        # Adding brewery name to the document
        brewery_word_document.add_paragraph(brewery_name, 'Title')

        # Adding brewery type
        brewery_word_document.add_paragraph('Brewery Type:', 'Heading 1')
        brewery_word_document.add_paragraph(brewery_type, 'Heading 3')

        # Adding brewery address
        brewery_word_document.add_paragraph('Brewery Address:', 'Heading 1')
        brewery_word_document.add_paragraph(brewery_address, 'Normal')
        brewery_word_document.add_paragraph(brewery_city, 'Normal')
        brewery_word_document.add_paragraph(brewery_state, 'Normal')

        # Add website URL to the document
        brewery_word_document.add_paragraph('Website', 'Heading 1')
        brewery_word_document.add_paragraph(brewery_website)


    else:
        print(f"Failed to retrieve details for brewery ID: {brewery_id}")

# # step 6
brewery_word_document.save('Brewery_Guide.docx')
