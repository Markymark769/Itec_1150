import json
import requests
import docx


def main():
    forecast = get_weather()

    # Access the key from the API that we want to locate, which is 'daily'
    list_of_weather_conditions = forecast['daily']

    # Create an empty list to store our data into from the 'daily' key from the dictionary
    weather = []

    # Loop over the list of weather conditions within API call
    for daily_data in list_of_weather_conditions:
        # Extracting the required data
        daily_weather = {
            'time': daily_data['time'],
            'temperature_2m_max': daily_data['temperature_2m_max'],
            'temperature_2m_min': daily_data['temperature_2m_min'],
            'daylight_duration': daily_data['daylight_duration'],
            'rain_sum': daily_data['rain_sum'],
            'showers_sum': daily_data['showers_sum'],
            'snowfall_sum': daily_data['snowfall_sum'],
            'wind_speed_10m_max': daily_data['wind_speed_10m_max']
        }
        weather.append(daily_weather)

    # Print the weather conditions for verification
    print(weather)

    # Create a Word document based on the weather data
    document = docx.Document()
    # Create a title
    document.add_paragraph('Weather for Picnic for Event', 'Title')

    # Add a header
    document.add_paragraph('By Mark Porraz', 'Heading 1')
    document.add_paragraph('Here is the data for time', 'Heading 2')

    # Loop through the weather data and add it to the document
    for index, daily_data in enumerate(weather, start=1):
        document.add_paragraph(f'Day {index}', 'Heading 3')
        document.add_paragraph(f'The times for the following days are:\n{daily_data["time"]}')

    # Save the document
    document.save('weather_for_picnic.docx')


def get_weather():
    # The URL for the weather API
    url = ('https://api.open-meteo.com/v1/forecast?latitude=52.52&longitude=13.41&daily=weather_code,temperature_2m_max,temperature_2m_min,sunrise,daylight_duration,rain_sum,showers_sum,snowfall_sum,wind_speed_10m_max,wind_gusts_10m_max,wind_direction_10m_dominant&timezone=America%2FChicago&start_date=2024-08-05&end_date=2024-08-18')
    try:
        response = requests.urlopen(url).read()
        data = json.loads(response)
        return data
    # except Exception as err:
    #     print(err)
    except:
        print('Sorry, unable to pull real data')
        return example_weather_data()



def example_weather_data():
    raw_weather_data = {
                            "latitude": 52.52,
                            "longitude": 13.419998,
                            "generationtime_ms": 1.114964485168457,
                            "utc_offset_seconds": -18000,
                            "timezone": "America/Chicago",
                            "timezone_abbreviation": "CDT",
                            "elevation": 38,
                            "current_units": {},
                            "current": {
                            "time": "2024-08-04T02:30",
                            "interval": 900,
                            "temperature_2m": 19,
                            "relative_humidity_2m": 90,
                            "apparent_temperature": 20.5,
                            "is_day": 1,
                            "precipitation": 0,
                            "rain": 0,
                            "showers": 0,
                            "snowfall": 0,
                            "weather_code": 61
                            },
                            "hourly_units": {},
                            "hourly": {
                            "time": [],
                            "temperature_2m": [],
                            "relative_humidity_2m": [],
                            "precipitation_probability": [],
                            "precipitation": [],
                            "rain": [],
                            "cloud_cover": [],
                            "cloud_cover_low": [],
                            "cloud_cover_mid": [],
                            "cloud_cover_high": [],
                            "wind_speed_10m": [],
                            "wind_speed_80m": [],
                            "wind_speed_120m": [],
                            "wind_speed_180m": [],
                            "temperature_120m": []
                            },
                            "daily_units": {},
                            "daily": {
                            "time": [],
                            "weather_code": [],
                            "temperature_2m_max": [],
                            "temperature_2m_min": [],
                            "sunrise": [],
                            "daylight_duration": [],
                            "rain_sum": [],
                            "showers_sum": [],
                            "snowfall_sum": [],
                            "wind_speed_10m_max": [
                            14.8,
                            7.9,
                            11.5,
                            9.6,
                            21.7,
                            13.4,
                            13.7,
                            21.9,
                            23.4,
                            15.8,
                            10.5,
                            20.9,
                            12.2,
                            21.3
                            ],
                            "wind_gusts_10m_max": [],
                            "wind_direction_10m_dominant": []
                            }
                            }
    return raw_weather_data


main()
