"""
Mark Porraz
weather API.py
this program makes a call to a weather API, places them into an excel file
next step is to place in function
"""
# import openpyxl
import requests
import docx


# The URL for the weather API
url = 'https://api.open-meteo.com/v1/forecast?latitude=52.52&longitude=13.41&current=temperature_2m,wind_speed_10m&hourly=temperature_2m,relative_humidity_2m,wind_speed_10m'
forecast = requests.get(url).json()  # correct syntax for getting a URL request in the JSON format

# accessing the key from the API that we want to locate, which is hours
list_of_weather_conditions = forecast['hourly']

# we are creating an empty list to store our data into from the 'hourly' key from the dictionary
weather = {}

# create a loop that goes over the list of weather conditions within API call
for hourly_data in list_of_weather_conditions:
    if weather:  # if the data in the hourly key exists, put it in an empty list called weather
        # below we are specific again on what data to pull out.
        time = hourly_data['time']  # we would like the hourly time
        temperature_2m = hourly_data['temperature_2m']  # we would like the hourly temperature
        relative_humidity_2m = hourly_data['relative_humidity_2m']  # we would like the hourly humidity
        wind_speed_10m = hourly_data['wind_speed_10m']  # we would like the hourly wind speed
    print(list_of_weather_conditions)  # print a new list of weather conditions

# set up for things in the docx

weather_word_document = docx.Document()
weather_word_document.add_paragraph('Current Weather Data', 'Title')
row = 4
for time, temperature_2m, relative_humidity_2m, wind_speed_10m in weather.items():
    weather_word_document.add_paragraph('The hourly data for time is the following', time)
    weather_word_document.add_paragraph('The hourly data for the temperature is the following', temperature_2m)
    weather_word_document.add_paragraph('The hourly data for humidity is the following', relative_humidity_2m)
    weather_word_document.add_paragraph('The hourly data for wind speed is the following', wind_speed_10m)

weather_word_document.save('Current Weather.docx')

# # setup spreadsheet stuff
# workbook = openpyxl.Workbook()
# worksheet = workbook.active
# worksheet.title = 'Weather API'
# worksheet.cell(1, 1, 'Time')
# worksheet.cell(1, 2, 'Temperature')
# worksheet.cell(1, 3, 'Humidity')
# worksheet.cell(1, 4, 'Wind Speed')
#
# worksheet.column_dimensions['A'].width = 40     # width for column A
# worksheet.column_dimensions['B'].width = 50     # width for column B
#
# workbook.save()
