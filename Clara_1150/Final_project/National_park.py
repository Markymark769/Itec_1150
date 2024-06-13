"""
Mark Porraz
National Park info from API
National Park Guide

1) Make a request from that list at random
2) Pick 5 parks from that list at random
3) make a Word document
4) for each of the random parks, make another API request for that park, write data
from the response to the Word document
5) (next week) get images
6) save the Word document


"""
# the goal is to make a complete working document based off this API
# the first thing we want to do with an API is look at what want to return
# then look to see if the information we have is accurate
# when looking at the API, the first website only gives us the name of the park, we neeed more info
# the other links have more info such as activities, contact info and descriptions
# from io import BytesIO
import requests
import random
import docx
from docx.shared import Inches
# from PIL import Image


# import urllib2
# import StringIO


# part 1
park_list_url = 'https://national-parks-1150.azurewebsites.net/api/list'
response_list_of_all_parks = requests.get(park_list_url).json()
# print(park_list_response)

# step 2
list_of_random_parks_chosen = random.sample(response_list_of_all_parks, 5)  # list, number of things which is park
# list url
print(list_of_random_parks_chosen)  # it is always a great idea to print to problem-solve

# step 3
park_word_document = docx.Document()

# step 4
for one_park in list_of_random_parks_chosen:
    print(one_park)
    park_code = one_park['park_code']
    print(park_code)


# https://national-parks-1150.azurewebsites.net/api/PARK CODE but replace with the park code
# for each of the random parks. make another API request to get details for that park
# write data from the response to word document.
    url_for_one_park_details = 'https://national-parks-1150.azurewebsites.net/api/' + park_code
    print(url_for_one_park_details)

    one_park_details_response = requests.get(url_for_one_park_details).json()
    print(one_park_details_response)


# step 5: get data from response and write to word document
# there is a video on downloading images
# TODO: comment code as you work
    # all text data except images
# park name
    park_name = one_park_details_response['name']
    print(park_name)

    park_word_document.add_paragraph(park_name, 'Title')

# description
    park_description = one_park_details_response['description']
    park_word_document.add_paragraph('Description', 'Heading 1')
    park_word_document.add_paragraph(park_description)

    # TODO: we need to figure out the weather, operating hours, and activites: this is going to be the same process as description
    # look for the weather overview and obtain the required text
# Weather Description
    park_weather_description = one_park_details_response['weather_overview']
    print(park_weather_description)
    park_word_document.add_paragraph('Weather Description', 'Heading 1')
    park_word_document.add_paragraph(park_weather_description)


    # how do we process the data in a list? We usually use a loop
    # for weather in park_weather_description:
    #     # TODO: use the bullet point style - it is in the videos/textbooks


    park_word_document.add_paragraph('Description', 'Heading 1')
    park_word_document.add_paragraph(park_description)
    # this is very similar to description and weather

# activities
    # note this is a bit different from the code before
    park_activities_list = one_park_details_response['activities']
    # # how do we process the data in a list? We usually use a loop
    for activity in park_activities_list:
    #     # TODO: use the bullet point style - it is in the videos/textbooks
        park_word_document.add_paragraph(activity)

    # # images (probably next time) --- a little more complex

#############   ISSUE HERE
##### Images #1

    # TODO: save the image from the JSON
    # TODO: maybe reference it as a variable
    # image_data = one_park_details_response['nps_park_images']
    # for image in image_data:
    #     image_url = image['url']
    #     caption = image.get('caption')
    #     credit = image.get('credit')
    #     credit_caption = caption + credit
    #     response = requests.get(image_url)
    #     image_content = response.content
    #     img = Image.open(BytesIO(image_content))
    #     park_word_document.add_paragraph('Picture','Heading 1')
    #     park_word_document.add_picture(BytesIO(image_content), width=docx.shared.Inches(4))
    #     park_word_document.add_paragraph(credit_caption,'Heading 2')
##### Images #2
    # get the list of the image info, use key from response
    image_data_list = one_park_details_response['nps_park_images']

    # It's a list so have to loop over each item. Each item is a dictionary
    for image_data in image_data_list:
        # image_data is a dictionary with title, caption, and URL for the image itself
        # for example,
        # {
        #     "altText": "a broad, red arch with rock pinnacles in the background",
        #     "caption": "Double O Arch is one of many large arches in the Devils Garden area",
        #     "credit": "NPS Photo",
        #     "title": "Double O Arch",
        #     "url": "https://www.nps.gov/common/uploads/structured_data/3C79931C-1DD8-B71B-0BF201E3DB540D04.jpg"
        # },
        image_title = image_data['title']  # this is a string, write it to the word doc
        image_download_url = image_data['url']  # this is the image download URL
        # download the image
        image_bytes = requests.get(image_download_url)  # no .json()
        # need a file name to store the image in
        park_image_filename = 'park_image.jpg'

        # Write the image to a file
        with open(park_image_filename, 'wb') as image_file:
            for chunk in image_bytes.iter_content():
                image_file.write(chunk)

        # Now the file is stored on your computer, can write to word doc
        # use the filename created above
        park_word_document.add_picture(park_image_filename, width=docx.shared.Inches(5))

#############   ISSUE HERE

# operating hours
    park_operating_hours = one_park_details_response['park_operating_hours']
    print(park_operating_hours)
    park_word_document.add_paragraph('Park Operating Hours', 'Heading 1')
    park_word_document.add_paragraph(park_operating_hours)

# contact info
    contact_info_dictionary = one_park_details_response['contact_info']
    print(contact_info_dictionary)
    # note: when we do this command
    print(one_park_details_response.keys())
    # displays this - these are the keys we can choose from
    # dict_keys(['activities', 'contact_info', 'description', 'location', 'name', 'nps_park_images',
    # 'park_code', 'park_operating_hours', 'weather_overview'])
    address_text = contact_info_dictionary['address']
    # note that /n
    print(address_text)

    # add header
    park_word_document.add_paragraph('Contact Information', 'Heading 1')
    # add sub-header
    park_word_document.add_paragraph('Address', 'Heading 2')
    # write the address next
    park_word_document.add_paragraph(address_text)

# add website
    park_word_document.add_paragraph('Website Info', 'Heading 1')
    park_word_document.add_paragraph(url_for_one_park_details)

# step 6

park_word_document.save('National Park Guide.docx')

# TODO: review the getting started video to know how to do this

# [] means a list - we are going to loop when used
# {} mean a dictionary - we use key and value as a pair

