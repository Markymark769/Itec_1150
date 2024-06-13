#############   ISSUE HERE
# attempt 1
    # TODO: save the image from the JSON
    # TODO: maybe reference it as a variable
    # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.run.Run.add_picture
    # park_word_document.add_paragraph('Images', 'Heading 1')
    # imageFile = park_word_document.add_picture(url_for_one_park_details, width=1.25, height=1.25)
    # park_word_document.add_paragraph(imageFile)
# attempt 2
    # park_word_document = Document()
    # url_for_one_park_details = 'park_image.jpeg'
    # park_word_document.add_paragraph('Images', 'Heading 1')
    # park_word_document.add_picture('park_image.jpeg', width=Inches(1.25), height=Inches(1.25))
    # park_word_document.save('document_with_image.docx')
# # attempt 3
#     park_word_document.add_paragraph('Images', 'Heading 1')
#     response = requests.get(url_for_one_park_details)  # no need to add stream=True
#     # Access the response body as bytes
#     #   then convert it to in-memory binary stream using `BytesIO`
#     binary_img = BytesIO(response.content)
#     park_word_document = Document()
#     # `add_picture` supports image path or stream, we use stream
#     park_word_document.add_picture(binary_img, width=Inches(2))
#     park_word_document.save('image.docx')
from docx.shared import Inches


# attempt 4
#     image_from_url = urllib2.urlopen(url_value)
#     io_url = StringIO.StringIO()
#     io_url.write(image_from_url.read())
#     io_url.seek(0)
#     try:
#         document.add_picture(io_url, width=Px(150))

# attempt 5
#     park_word_document = Document()
#     url_for_one_park_details = 'park_image.jpeg'
#     park_word_document.add_picture('park_image.jpg', width=Inches(1.25), height=Inches(1.25))
#
#     image_from_url = urllib2.urlopen('https://national-parks-1150.azurewebsites.net/api/' + park_code)
#     image_response = requests.get(image_from_url).json()
#     io_url = StringIO.StringIO()
#     io_url.write(image_from_url.read())
#     io_url.seek(0)
#     try:
#         park_word_document.add_picture(io_url, width=Px(150))
# attempt 6
    image_from_url = urllib2.urlopen('https://national-parks-1150.azurewebsites.net/api/' + park_code)
    image_response = requests.get(image_from_url).json()
    io_url = StringIO.StringIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)

    # Create a new Word document
    park_word_document = Document()

    try:
        park_word_document.add_picture(io_url, width=Inches(1.5))  # Use Inches or Cm instead of Px
    except Exception as e:
        print(f"Error adding picture: {e}")

    # Save the document
    park_word_document.save('document_with_image.docx')
# attempt 7
# park_word_document = docx.Document()
    # url_for_one_park_details = 'park_image.jpeg'
    # park_word_document.add_picture('park_image.jpg', width=Inches(1.25), height=Inches(1.25))
    # #
    # image_from_url = urlopen('https://national-parks-1150.azurewebsites.net/api/OLYM')
    #############   ISSUE HERE

# get the list of the image info, use key from response
#

# It's a list so have to loop over each item. Each item is a dictionary
#
    # image_data is a dictionary with title, caption, and URL for the image itself
    # for example,
    #  {
    #      "altText": "a broad, red arch with rock pinnacles in the background",
    #      "caption": "Double O Arch is one of many large arches in the Devils Garden area",
    #      "credit": "NPS Photo",
    #      "title": "Double O Arch",
    #      "url": "https://www.nps.gov/common/uploads/structured_data/3C79931C-1DD8-B71B-0BF201E3DB540D04.jpg"
    #  },
    # image_title = image_data['title']  # this is a string, write it to the word doc
    # image_download_url = image_data['url']  # this is the image download URL
    # # download the image
    # image_bytes = requests.get(image_download_url)  # no .json()
    # # need a file name to store the image in
    # park_image_filename = 'park_image.jpg'
    #
    # # Write the image to a file
    # with open(park_image_filename, 'wb') as image_file:
    #     for chunk in image_bytes.iter_content():
    #         image_file.write(chunk)
    #
    # # Now the file is stored on your computer, can write to word doc
    # # use the filename created above
    # park_word_document.add_picture(park_image_filename, width=docx.shared.Inches(5))
# solution 1:
#     image_data = one_park_details_response['nps_park_images']
#     for image in image_data:
#         image_url = image['url']
#         caption = image.get('caption')
#         credit = image.get('credit')
#         credit_caption = caption + credit
#         response = requests.get(image_url)
#         image_content = response.content
#         img = Image.open(BytesIO(image_content))
#         park_word_document.add_paragraph('Picture','Heading 1')
#         park_word_document.add_picture(BytesIO(image_content), width=docx.shared.Inches(4))
#         park_word_document.add_paragraph(credit_caption,'Heading 2')
# solution 2